"""Agent 1 简历分析 API 接口测试。

覆盖两个接口：
- POST /api/v1/resume/upload — 上传简历，返回 task_id 和 file_type
- GET  /api/v1/resume/{task_id}/analysis — 获取结构化简历和评估结果
"""

from pathlib import Path
import asyncio
import sys

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.api.v1.router import api_router  # noqa: E402
from app.api.deps import get_current_user  # noqa: E402
from app.core.config import settings  # noqa: E402
from app.services.resume_tasks import (  # noqa: E402
    analyze_resume_task,
    resume_task_store,
)


def _create_app() -> FastAPI:
    """创建挂载了 v1 路由的 FastAPI 应用实例。"""
    app = FastAPI()
    app.dependency_overrides[get_current_user] = lambda: {
        "id": "test-user",
        "email": "test@example.com",
        "username": "tester",
    }
    app.include_router(api_router, prefix="/api/v1")
    return app


def _create_app_client() -> TestClient:
    return TestClient(_create_app())


def _make_docx(path: Path) -> None:
    """创建包含简历内容的 DOCX 文件。"""
    document = Document()
    document.add_paragraph("姓名：王五")
    document.add_paragraph("电话：13700137000")
    document.add_paragraph("邮箱：wangwu@example.com")
    document.add_paragraph("技能：Python FastAPI Vue")
    document.add_paragraph("项目经历：简历分析系统")
    document.save(path)


def _make_pdf(path: Path) -> None:
    """创建一个带 %PDF 文件头的最小 PDF 文件。"""
    path.write_bytes(b"%PDF-1.7\n1 0 obj\n<< /Type /Catalog >>\nendobj\n")


# ========================================================================
# POST /api/v1/resume/upload
# ========================================================================


class TestUploadResume:
    """POST /api/v1/resume/upload 接口测试。"""

    def setup_method(self):
        """每个测试前清空 task store，避免状态污染。"""
        resume_task_store._tasks.clear()
        settings.AI_ANALYSIS_ENABLED = False

    def test_upload_docx_returns_task_id_and_file_type(self, tmp_path):
        """上传 docx 返回 task_id、file_type 和 current_stage。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            response = client.post(
                "/api/v1/resume/upload",
                files={
                    "file": (
                        "resume.docx",
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert response.status_code == 200
        body = response.json()
        assert "task_id" in body
        assert body["file_type"] == "docx"
        assert body["current_stage"] == "upload"

    def test_upload_defers_analysis_until_needed(self, tmp_path):
        """上传阶段只保存文件，不提前消耗一次 LLM 调用。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={
                    "file": (
                        "resume.docx",
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        task_id = upload_resp.json()["task_id"]

        analysis = client.get(f"/api/v1/resume/{task_id}/analysis")
        assert analysis.status_code == 200
        assert analysis.json()["current_stage"] == "upload"
        assert analysis.json()["structured_resume"] == {}

    def test_upload_response_schema(self, tmp_path):
        """上传响应包含 task_id、file_type、current_stage 三个字段。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            response = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        body = response.json()
        assert set(body.keys()) == {"task_id", "file_type", "current_stage"}

    def test_upload_rejects_unsupported_file(self):
        response = _create_app_client().post(
            "/api/v1/resume/upload",
            files={"file": ("resume.txt", b"not a resume", "text/plain")},
        )

        assert response.status_code == 400
        assert "PDF" in response.json()["detail"]

    def test_upload_rejects_oversized_file(self, monkeypatch):
        monkeypatch.setattr(settings, "RESUME_MAX_UPLOAD_BYTES", 8)
        response = _create_app_client().post(
            "/api/v1/resume/upload",
            files={"file": ("resume.pdf", b"%PDF-1.7 oversized", "application/pdf")},
        )

        assert response.status_code == 413


# ========================================================================
# GET /api/v1/resume/{task_id}/analysis
# ========================================================================


class TestGetResumeAnalysis:
    """GET /api/v1/resume/{task_id}/analysis 接口测试。"""

    def setup_method(self):
        """每个测试前清空 task store。"""
        resume_task_store._tasks.clear()
        settings.AI_ANALYSIS_ENABLED = False

    def test_get_analysis_success(self, tmp_path):
        """成功获取已完成的分析结果。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        task_id = upload_resp.json()["task_id"]
        asyncio.run(analyze_resume_task(task_id))
        response = client.get(f"/api/v1/resume/{task_id}/analysis")

        assert response.status_code == 200
        data = response.json()
        assert data["task_id"] == task_id
        assert data["current_stage"] == "done"
        assert data["error"] is None
        assert data["file_type"] == "docx"

    def test_get_analysis_structured_resume_content(self, tmp_path):
        """分析结果中 structured_resume 包含正确提取的信息。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        task_id = upload_resp.json()["task_id"]
        asyncio.run(analyze_resume_task(task_id))
        data = client.get(f"/api/v1/resume/{task_id}/analysis").json()

        sr = data["structured_resume"]
        assert sr["basic_info"]["phone"] == "13700137000"
        assert sr["basic_info"]["email"] == "wangwu@example.com"
        assert sr["basic_info"]["name"] == "王五"
        assert "Python" in sr["skills"]

    def test_get_analysis_evaluation_content(self, tmp_path):
        """分析结果中 evaluation 包含完整度评分和优势/劣势。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        task_id = upload_resp.json()["task_id"]
        asyncio.run(analyze_resume_task(task_id))
        data = client.get(f"/api/v1/resume/{task_id}/analysis").json()

        ev = data["evaluation"]
        assert "completeness_score" in ev
        assert 0 <= ev["completeness_score"] <= 100
        assert isinstance(ev["strengths"], list)
        assert isinstance(ev["weaknesses"], list)
        assert isinstance(ev["missing_sections"], list)

    def test_get_analysis_response_schema(self, tmp_path):
        """分析响应包含所有预定义字段。"""
        app = _create_app()
        client = TestClient(app)

        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        task_id = upload_resp.json()["task_id"]
        data = client.get(f"/api/v1/resume/{task_id}/analysis").json()

        assert set(data.keys()) == {
            "task_id",
            "current_stage",
            "error",
            "file_type",
            "converted_file_path",
            "structured_resume",
            "evaluation",
            "job_search_results",
            "selected_jd_url",
        }

    def test_get_analysis_not_found(self):
        """不存在的 task_id 返回 404。"""
        app = _create_app()
        client = TestClient(app)

        response = client.get("/api/v1/resume/nonexistent-task-id/analysis")
        assert response.status_code == 404
        assert "not found" in response.json()["detail"].lower()

    def test_get_analysis_hides_other_users_task(self, tmp_path):
        app = _create_app()
        client = TestClient(app)
        docx = tmp_path / "resume.docx"
        _make_docx(docx)
        with docx.open("rb") as file:
            task_id = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", file)},
            ).json()["task_id"]

        app.dependency_overrides[get_current_user] = lambda: {
            "id": "another-user",
            "email": "other@example.com",
            "username": "other",
        }

        response = client.get(f"/api/v1/resume/{task_id}/analysis")
        assert response.status_code == 404


# ========================================================================
# 端到端流程测试
# ========================================================================


class TestEndToEndFlow:
    """端到端流程测试：上传 → 分析 → 获取结果。"""

    def setup_method(self):
        resume_task_store._tasks.clear()
        settings.AI_ANALYSIS_ENABLED = False

    def test_full_upload_analyze_retrieve_flow(self, tmp_path):
        """完整的上传 → 分析 → 获取结果流程。"""
        app = _create_app()
        client = TestClient(app)

        # 1. 上传
        docx = tmp_path / "resume.docx"
        _make_docx(docx)

        with docx.open("rb") as f:
            upload_resp = client.post(
                "/api/v1/resume/upload",
                files={"file": ("resume.docx", f)},
            )

        assert upload_resp.status_code == 200
        task_id = upload_resp.json()["task_id"]

        # 2. 由后续业务动作按需执行 Agent 1
        asyncio.run(analyze_resume_task(task_id))
        analysis_resp = client.get(f"/api/v1/resume/{task_id}/analysis")
        assert analysis_resp.status_code == 200

        data = analysis_resp.json()
        # 验证完整流程产出
        assert data["current_stage"] == "done"
        assert data["file_type"] == "docx"
        assert data["structured_resume"]["basic_info"]["name"] == "王五"
        assert data["structured_resume"]["basic_info"]["phone"] == "13700137000"
        assert data["evaluation"]["completeness_score"] > 0

    def test_multiple_uploads_isolated(self, tmp_path):
        """多次上传互不干扰，各自返回独立 task_id。"""
        app = _create_app()
        client = TestClient(app)

        task_ids = []
        for i in range(3):
            docx = tmp_path / f"resume_{i}.docx"
            document = Document()
            document.add_paragraph(f"姓名：测试用户{i}")
            document.add_paragraph(f"电话：1380013800{i}")
            document.save(docx)

            with docx.open("rb") as f:
                resp = client.post(
                    "/api/v1/resume/upload",
                    files={"file": (f"resume_{i}.docx", f)},
                )
            assert resp.status_code == 200
            task_ids.append(resp.json()["task_id"])

        # 三个 task_id 互不相同
        assert len(set(task_ids)) == 3

        # 每个都能独立获取分析结果
        for i, tid in enumerate(task_ids):
            asyncio.run(analyze_resume_task(tid))
            data = client.get(f"/api/v1/resume/{tid}/analysis").json()
            assert data["structured_resume"]["basic_info"]["phone"] == f"1380013800{i}"


if __name__ == "__main__":
    import pytest

    raise SystemExit(pytest.main([__file__, "-v"]))
