"""Agent 1 简历分析 Agent 单元测试。

覆盖四个 Tool（file_type_detector / pdf_to_word_converter /
document_text_extractor / resume_structurer）以及 Agent 的线性执行流程
和 State 传递逻辑。
"""

from pathlib import Path
import sys
import zipfile

from docx import Document

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.agent.resume_analysis_agent import (  # noqa: E402
    ResumeAnalysisAgent,
    convert_pdf_node,
    detect_file_type_node,
    extract_text_node,
    run_resume_analysis_agent,
    structure_resume_node,
)
from app.schema.workflow_state import initial_workflow_state  # noqa: E402
from app.tools.document_text_extractor import document_text_extractor  # noqa: E402
from app.tools.file_type_detector import file_type_detector  # noqa: E402
from app.tools.pdf_to_word_converter import pdf_to_word_converter  # noqa: E402
from app.tools.resume_structurer import (  # noqa: E402
    build_resume_structure_prompt,
    resume_structurer,
)


# ===== 辅助函数 =====


def make_docx(path: Path, lines: list[str]) -> None:
    """创建一个包含指定段落的 DOCX 文件。"""
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(path)


def make_pdf(path: Path) -> None:
    """创建一个最简单的 PDF 文件（带 %PDF 头）。"""
    path.write_bytes(b"%PDF-1.7\n1 0 obj\n<< /Type /Catalog >>\nendobj\n")


# ========================================================================
# Tool 1.1: file_type_detector
# ========================================================================


class TestFileTypeDetector:
    """Tool 1.1: 文件类型检测器测试。"""

    def test_detect_pdf_by_header(self, tmp_path):
        """通过 %PDF 文件头识别 PDF。"""
        pdf = tmp_path / "resume.pdf"
        pdf.write_bytes(b"%PDF-1.7\nbody")
        result = file_type_detector(str(pdf))
        assert result["file_type"] == "pdf"
        assert result["file_path"] == str(pdf)

    def test_detect_docx_by_header(self, tmp_path):
        """通过 PK(OOXML) 文件头 + .docx 扩展名识别 DOCX。"""
        docx = tmp_path / "resume.docx"
        docx.write_bytes(b"PK\x03\x04body")
        result = file_type_detector(str(docx))
        assert result["file_type"] == "docx"

    def test_detect_doc_by_header(self, tmp_path):
        """通过 OLE2 文件头识别旧版 DOC。"""
        doc = tmp_path / "resume.doc"
        doc.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1body")
        result = file_type_detector(str(doc))
        assert result["file_type"] == "doc"

    def test_detect_unknown_type(self, tmp_path):
        """无法识别的文件类型返回 unknown。"""
        txt = tmp_path / "resume.txt"
        txt.write_text("hello world", encoding="utf-8")
        result = file_type_detector(str(txt))
        assert result["file_type"] == "unknown"

    def test_detect_nonexistent_file(self, tmp_path):
        """文件不存在时按扩展名兜底，txt 返回 unknown。"""
        result = file_type_detector(str(tmp_path / "missing.pdf"))
        assert result["file_type"] == "pdf"  # 扩展名兜底

    def test_detect_nonexistent_unknown_extension(self, tmp_path):
        """文件不存在且扩展名未知时返回 unknown。"""
        result = file_type_detector(str(tmp_path / "missing.xyz"))
        assert result["file_type"] == "unknown"

    def test_output_schema(self, tmp_path):
        """输出包含 file_type 和 file_path 两个键。"""
        pdf = tmp_path / "resume.pdf"
        pdf.write_bytes(b"%PDF-1.7")
        result = file_type_detector(str(pdf))
        assert set(result.keys()) == {"file_type", "file_path"}


# ========================================================================
# Tool 1.2: pdf_to_word_converter
# ========================================================================


class TestPdfToWordConverter:
    """Tool 1.2: PDF 转 Word 转换器测试。"""

    def test_skip_docx_input(self, tmp_path):
        """输入是 DOCX 时直接跳过，返回成功。"""
        docx = tmp_path / "resume.docx"
        docx.write_bytes(b"PK\x03\x04body")
        result = pdf_to_word_converter(str(docx))
        assert result == {"converted_path": str(docx), "success": True}

    def test_unsupported_file_type(self, tmp_path):
        """非 PDF/DOCX 文件返回 success=False。"""
        txt = tmp_path / "resume.txt"
        txt.write_text("hello", encoding="utf-8")
        result = pdf_to_word_converter(str(txt))
        assert result["success"] is False

    def test_pdf_conversion_output_schema(self, tmp_path):
        """PDF 转换结果包含 converted_path 和 success 两个键。"""
        pdf = tmp_path / "resume.pdf"
        make_pdf(pdf)
        result = pdf_to_word_converter(str(pdf))
        assert "converted_path" in result
        assert "success" in result
        # 转换可能成功也可能失败（取决于 pdf2docx 是否安装及文件是否有效）
        # 但路径应该以 .docx 结尾
        assert result["converted_path"].endswith(".docx")

    def test_pdf_pages_are_embedded_without_reflow(self, tmp_path):
        """PDF 页面应以原始渲染像素嵌入 Word，不能重排文字或模板。"""
        import fitz

        pdf_path = tmp_path / "resume.pdf"
        pdf = fitz.open()
        for label in ("PAGE ONE", "PAGE TWO"):
            page = pdf.new_page(width=595, height=842)
            page.insert_text((72, 72), label, fontsize=18)
        pdf.save(pdf_path)
        pdf.close()

        with fitz.open(pdf_path) as source:
            expected_pages = [
                page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False).tobytes("png")
                for page in source
            ]

        result = pdf_to_word_converter(str(pdf_path))

        assert result["success"] is True
        with zipfile.ZipFile(result["converted_path"]) as package:
            image_names = sorted(
                name for name in package.namelist() if name.startswith("word/media/")
            )
            embedded_pages = [package.read(name) for name in image_names]
        assert embedded_pages == expected_pages


# ========================================================================
# Tool 1.3: document_text_extractor
# ========================================================================


class TestDocumentTextExtractor:
    """Tool 1.3: 文档文本提取器测试。"""

    def test_extract_docx_text(self, tmp_path):
        """从 DOCX 提取段落文本。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["姓名：张三", "技能：Python, FastAPI"])
        result = document_text_extractor(str(docx))
        assert "张三" in result["raw_text"]
        assert "FastAPI" in result["raw_text"]
        assert result["char_count"] == len(result["raw_text"])

    def test_extract_docx_includes_tables(self, tmp_path):
        """从 DOCX 表格中提取文本。"""
        doc = Document()
        doc.add_paragraph("标题")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "公司"
        table.cell(0, 1).text = "职位"
        table.cell(1, 0).text = "ABC公司"
        table.cell(1, 1).text = "后端开发"
        docx = tmp_path / "resume.docx"
        doc.save(docx)

        result = document_text_extractor(str(docx))
        assert "标题" in result["raw_text"]
        assert "ABC公司" in result["raw_text"]
        assert "后端开发" in result["raw_text"]

    def test_extract_unsupported_type(self, tmp_path):
        """不支持的文件类型返回空文本。"""
        txt = tmp_path / "resume.txt"
        txt.write_text("hello", encoding="utf-8")
        result = document_text_extractor(str(txt))
        assert result["raw_text"] == ""
        assert result["char_count"] == 0

    def test_output_schema(self, tmp_path):
        """输出包含 raw_text 和 char_count 两个键。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["test"])
        result = document_text_extractor(str(docx))
        assert set(result.keys()) == {"raw_text", "char_count"}


# ========================================================================
# Tool 1.4: resume_structurer
# ========================================================================


class TestResumeStructurer:
    """Tool 1.4: 简历结构化提取 + 评估测试。"""

    def test_output_schema_keys(self):
        """输出包含 structured_resume 和 evaluation 两个顶层键。"""
        result = resume_structurer("some text")
        assert set(result.keys()) == {"structured_resume", "evaluation"}

    def test_structured_resume_schema(self):
        """structured_resume 包含所有预定义字段。"""
        result = resume_structurer("some text")
        sr = result["structured_resume"]
        assert set(sr.keys()) == {
            "basic_info",
            "education",
            "work_experience",
            "project_experience",
            "skills",
            "self_evaluation",
        }
        assert set(sr["basic_info"].keys()) == {"name", "phone", "email", "location"}

    def test_evaluation_schema(self):
        """evaluation 包含所有预定义字段。"""
        result = resume_structurer("some text")
        ev = result["evaluation"]
        assert set(ev.keys()) == {
            "analysis_source",
            "completeness_score",
            "overall_summary",
            "strengths",
            "weaknesses",
            "missing_sections",
            "section_reviews",
            "improvement_suggestions",
            "keyword_analysis",
            "ats_readability",
            "risk_points",
            "rewrite_examples",
            "llm_error",
        }
        assert set(ev["keyword_analysis"].keys()) == {
            "detected_keywords",
            "missing_keywords",
            "keyword_density_comment",
        }

    def test_completeness_score_range(self):
        """完整度评分在 0-100 范围内。"""
        result = resume_structurer("some text")
        assert 0 <= result["evaluation"]["completeness_score"] <= 100

    def test_fallback_extracts_phone_and_email(self):
        """fallback 模式提取电话和邮箱。"""
        text = (
            "姓名：张三\n电话：13800138000\n邮箱：test@example.com\n技能：Python、SQL"
        )
        result = resume_structurer(text, llm=None)
        assert result["structured_resume"]["basic_info"]["phone"] == "13800138000"
        assert result["structured_resume"]["basic_info"]["email"] == "test@example.com"
        assert result["structured_resume"]["basic_info"]["name"] == "张三"
        assert "Python" in result["structured_resume"]["skills"]

    def test_fallback_extracts_explicit_technologies_without_skills_heading(self):
        """没有技能标题时，从项目正文提取明确出现过的技术词。"""
        text = """
        项目经历
        DocMind —— 基于 RAG 的智能知识库问答系统
        使用 FAISS 构建向量索引，通过 Redis 实现缓存和限流，使用 Docker 部署。
        """

        result = resume_structurer(text, llm=None)

        assert {"RAG", "FAISS", "Redis", "Docker"}.issubset(
            set(result["structured_resume"]["skills"])
        )

    def test_fallback_extracts_pipe_style_pdf_resume(self):
        """fallback 能处理 PDF 提取出的管道分隔中文简历信息。"""
        text = """
        个人优势
        熟悉 Python、FastAPI 和常见智能应用开发流程，具备团队协作经验。
        教育经历
        示例大学 | 本科 | 计算机科学与技术 | 2020-2024
        项目经历
        智能客服系统
        1. 技术栈
        基于 Python 3.10+，使用 LangChain 构建 Agent，通过 WebSocket 处理实时消息。
        专业技能
        1. 掌握Python 编程，熟练掌握 LangChain 框架及 ReAct 模式。
        2. 熟悉 WebSocket 协议与异步编程模型。
        李明 | 李明 | 李明
        男 | 年龄：24岁 | | 13800138000 | | candidate@example.com
        求职意向：软件开发 | 期望薪资：10-15K | 期望城市：北京
        智能客服系统 | 开发工程师 | 2023.01-2023.12
        """
        result = resume_structurer(text, llm=None)
        resume = result["structured_resume"]

        assert resume["basic_info"]["name"] == "李明"
        assert resume["basic_info"]["phone"] == "13800138000"
        assert resume["basic_info"]["email"] == "candidate@example.com"
        assert resume["basic_info"]["location"] == "北京"
        assert resume["education"][0]["school"] == "示例大学"
        assert resume["education"][0]["major"] == "计算机科学与技术"
        assert resume["education"][0]["degree"] == "本科"
        assert resume["education"][0]["period"] == "2020-2024"
        assert resume["project_experience"][0]["name"] == "智能客服系统"
        assert resume["project_experience"][0]["role"] == "开发工程师"
        assert resume["skills"]
        assert result["evaluation"]["completeness_score"] >= 80

    def test_fallback_completeness_score(self):
        """fallback 模式包含多种信息时评分较高。"""
        text = (
            "姓名：张三\n电话：13800138000\n邮箱：test@example.com\n"
            "教育：中山大学 本科\n工作经历：某公司\n项目经历：数据平台\n技能：Python"
        )
        result = resume_structurer(text, llm=None)
        assert result["evaluation"]["completeness_score"] == 100

    def test_fallback_missing_sections(self):
        """fallback 模式检测缺失的部分。"""
        text = "姓名：张三\n电话：13800138000"
        result = resume_structurer(text, llm=None)
        missing = result["evaluation"]["missing_sections"]
        assert "教育经历" in missing
        assert "工作经历" in missing
        assert "项目经历" in missing
        assert "技能清单" in missing

    def test_injected_llm_dict_response(self):
        """注入 LLM 返回 dict 时正确解析。"""

        def fake_llm(prompt: str):
            assert "简历原文" in prompt
            return {
                "structured_resume": {
                    "basic_info": {
                        "name": "张三",
                        "phone": "13800138000",
                        "email": "a@example.com",
                        "location": "广州",
                    },
                    "education": [
                        {
                            "school": "中山大学",
                            "major": "软件工程",
                            "degree": "本科",
                            "period": "2020-2024",
                        }
                    ],
                    "work_experience": [],
                    "project_experience": [
                        {"name": "招聘助手", "role": "后端", "description": "FastAPI"}
                    ],
                    "skills": ["Python", "FastAPI"],
                    "self_evaluation": "学习能力强",
                },
                "evaluation": {
                    "completeness_score": 88,
                    "strengths": ["项目清晰"],
                    "weaknesses": ["工作经历较少"],
                    "missing_sections": ["工作经历"],
                },
            }

        result = resume_structurer("姓名：张三", llm=fake_llm)
        assert result["structured_resume"]["basic_info"]["name"] == "张三"
        assert result["evaluation"]["completeness_score"] == 88
        assert "项目清晰" in result["evaluation"]["strengths"]

    def test_injected_llm_json_string_response(self):
        """注入 LLM 返回 JSON 字符串时正确解析。"""
        import json

        def fake_llm(prompt: str):
            return json.dumps(
                {
                    "structured_resume": {
                        "basic_info": {
                            "name": "李四",
                            "phone": "",
                            "email": "",
                            "location": "",
                        },
                        "education": [],
                        "work_experience": [],
                        "project_experience": [],
                        "skills": ["Java"],
                        "self_evaluation": "",
                    },
                    "evaluation": {
                        "completeness_score": 50,
                        "strengths": [],
                        "weaknesses": ["缺少经历"],
                        "missing_sections": ["教育经历", "工作经历"],
                    },
                }
            )

        result = resume_structurer("姓名：李四", llm=fake_llm)
        assert result["structured_resume"]["basic_info"]["name"] == "李四"
        assert result["evaluation"]["completeness_score"] == 50

    def test_injected_llm_markdown_wrapped_response(self):
        """注入 LLM 返回 markdown 代码块包裹的 JSON 时正确解析。"""

        def fake_llm(prompt: str):
            return '```json\n{"structured_resume": {"basic_info": {"name": "王五", "phone": "", "email": "", "location": ""}, "education": [], "work_experience": [], "project_experience": [], "skills": [], "self_evaluation": ""}, "evaluation": {"completeness_score": 20, "strengths": [], "weaknesses": [], "missing_sections": []}}\n```'

        result = resume_structurer("姓名：王五", llm=fake_llm)
        assert result["structured_resume"]["basic_info"]["name"] == "王五"

    def test_injected_llm_score_clamped(self):
        """LLM 返回的评分超出范围时被 clamp 到 0-100。"""

        def fake_llm(prompt: str):
            return {
                "structured_resume": {
                    "basic_info": {
                        "name": "",
                        "phone": "",
                        "email": "",
                        "location": "",
                    },
                    "education": [],
                    "work_experience": [],
                    "project_experience": [],
                    "skills": [],
                    "self_evaluation": "",
                },
                "evaluation": {
                    "completeness_score": 150,
                    "strengths": [],
                    "weaknesses": [],
                    "missing_sections": [],
                },
            }

        result = resume_structurer("text", llm=fake_llm)
        assert result["evaluation"]["completeness_score"] == 100

    def test_configured_llm_path_uses_model_adapter(self, monkeypatch):
        """启用已配置模型时应调用模型适配器，而不是发生名称解析错误。"""

        def fake_chat_completion(prompt: str, system_prompt: str = ""):
            assert "姓名：赵六" in prompt
            assert system_prompt
            return {
                "structured_resume": {
                    "basic_info": {
                        "name": "赵六",
                        "phone": "",
                        "email": "",
                        "location": "",
                    },
                    "education": [],
                    "work_experience": [],
                    "project_experience": [],
                    "skills": [],
                    "self_evaluation": "",
                },
                "evaluation": {"analysis_source": "llm", "completeness_score": 20},
            }

        monkeypatch.setattr(
            "app.model.openai_compatible.chat_completion", fake_chat_completion
        )

        result = resume_structurer("姓名：赵六", use_configured_llm=True)

        assert result["structured_resume"]["basic_info"]["name"] == "赵六"
        assert result["evaluation"]["analysis_source"] == "llm"
        assert result["evaluation"]["llm_error"] == ""

    def test_build_prompt_contains_schema(self):
        """build_resume_structure_prompt 包含 schema 定义和原文。"""
        prompt = build_resume_structure_prompt("测试原文")
        assert "structured_resume" in prompt
        assert "basic_info" in prompt
        assert "evaluation" in prompt
        assert "completeness_score" in prompt
        assert "测试原文" in prompt


# ========================================================================
# Agent: 线性执行 + State 传递
# ========================================================================


class TestResumeAnalysisAgentNodes:
    """Agent 各节点函数单独测试。"""

    def test_detect_file_type_node(self, tmp_path):
        """detect_file_type_node 写入 state["file_type"]。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["test"])
        state = initial_workflow_state("task-1", str(docx))

        detect_file_type_node(state)

        assert state["file_type"] == "docx"

    def test_convert_pdf_node_skips_docx(self, tmp_path):
        """convert_pdf_node 对 docx 设置 converted_file_path = 原路径。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["test"])
        state = initial_workflow_state("task-1", str(docx))
        state["file_type"] = "docx"

        convert_pdf_node(state)

        assert state["converted_file_path"] == str(docx)

    def test_convert_pdf_node_unknown_type(self):
        """convert_pdf_node 对 unknown 类型设置 converted_file_path = None。"""
        state = initial_workflow_state("task-1", "/fake/path.txt")
        state["file_type"] = "unknown"

        convert_pdf_node(state)

        assert state["converted_file_path"] is None

    def test_extract_text_node(self, tmp_path):
        """extract_text_node 写入 state["raw_text"]。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["姓名：张三", "电话：13800138000"])
        state = initial_workflow_state("task-1", str(docx))
        state["file_type"] = "docx"
        state["converted_file_path"] = str(docx)

        extract_text_node(state)

        assert "张三" in state["raw_text"]
        assert "13800138000" in state["raw_text"]

    def test_extract_text_node_reads_original_pdf(self, tmp_path, monkeypatch):
        """PDF 分析必须读取原文件，不能读取为了展示生成的图片版 Word。"""
        pdf_path = tmp_path / "resume.pdf"
        converted_path = tmp_path / "resume.docx"
        pdf_path.write_bytes(b"%PDF-original")
        converted_path.write_bytes(b"PK\x03\x04converted")
        requested_paths = []

        def fake_extractor(file_path: str):
            requested_paths.append(file_path)
            return {"raw_text": "原始 PDF 文本", "char_count": 9}

        monkeypatch.setattr(
            "app.agent.resume_analysis_agent.document_text_extractor",
            fake_extractor,
        )
        state = initial_workflow_state("task-pdf", str(pdf_path))
        state["file_type"] = "pdf"
        state["converted_file_path"] = str(converted_path)

        extract_text_node(state)

        assert requested_paths == [str(pdf_path)]
        assert state["raw_text"] == "原始 PDF 文本"

    def test_structure_resume_node(self):
        """structure_resume_node 写入 structured_resume 和 resume_evaluation。"""
        state = initial_workflow_state("task-1", "/fake/path.docx")
        state["raw_text"] = "姓名：张三\n电话：13800138000\n技能：Python"

        structure_resume_node(state, use_configured_llm=False)

        assert state["structured_resume"]["basic_info"]["phone"] == "13800138000"
        assert state["structured_resume"]["basic_info"]["name"] == "张三"
        assert state["resume_evaluation"]["completeness_score"] > 0


class TestResumeAnalysisAgentRun:
    """Agent 完整线性执行测试。"""

    def test_run_docx_full_pipeline(self, tmp_path):
        """docx 文件的完整流程：1.1 → 1.2(skip) → 1.3 → 1.4。"""
        docx = tmp_path / "resume.docx"
        make_docx(
            docx,
            [
                "姓名：李四",
                "电话：13900139000",
                "技能：Python、SQL",
                "项目经历：数据平台",
            ],
        )
        state = initial_workflow_state(task_id="task-1", file_path=str(docx))

        final_state = run_resume_analysis_agent(state, use_configured_llm=False)

        assert final_state["current_stage"] == "done"
        assert final_state["error"] is None
        assert final_state["file_type"] == "docx"
        # Tool 1.2 跳过转换，converted_file_path = 原路径
        assert final_state["converted_file_path"] == str(docx)
        # Tool 1.3 提取文本
        assert "李四" in final_state["raw_text"]
        # Tool 1.4 结构化 + 评估
        assert final_state["structured_resume"]["basic_info"]["phone"] == "13900139000"
        assert final_state["resume_evaluation"]["completeness_score"] > 0

    def test_run_unknown_file_type_sets_error(self, tmp_path):
        """不支持的文件类型设置 error 状态。"""
        txt = tmp_path / "resume.txt"
        txt.write_text("hello", encoding="utf-8")
        state = initial_workflow_state(task_id="task-1", file_path=str(txt))

        final_state = run_resume_analysis_agent(state, use_configured_llm=False)

        assert final_state["current_stage"] == "error"
        assert "Unsupported" in final_state["error"]

    def test_run_doc_file_type_sets_error(self, tmp_path):
        """.doc 文件类型设置 error 状态。"""
        doc = tmp_path / "resume.doc"
        doc.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1body")
        state = initial_workflow_state(task_id="task-1", file_path=str(doc))

        final_state = run_resume_analysis_agent(state, use_configured_llm=False)

        assert final_state["current_stage"] == "error"
        assert ".doc" in final_state["error"]

    def test_state_callback_invoked(self, tmp_path):
        """on_state_update 回调在每个阶段被调用。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["姓名：张三", "技能：Python"])
        state = initial_workflow_state(task_id="task-1", file_path=str(docx))

        updates: list[str] = []

        def callback(s):
            updates.append(s["current_stage"])

        run_resume_analysis_agent(
            state, on_state_update=callback, use_configured_llm=False
        )

        # 应该至少有 analyzing 和 done
        assert "analyzing" in updates
        assert "done" in updates

    def test_run_with_injected_llm(self, tmp_path):
        """注入 LLM 时使用 LLM 结果而非 fallback。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["姓名：赵六", "电话：13700137000"])
        state = initial_workflow_state(task_id="task-1", file_path=str(docx))

        def fake_llm(prompt: str):
            return {
                "structured_resume": {
                    "basic_info": {
                        "name": "赵六",
                        "phone": "13700137000",
                        "email": "zhao@example.com",
                        "location": "深圳",
                    },
                    "education": [],
                    "work_experience": [],
                    "project_experience": [],
                    "skills": ["Go"],
                    "self_evaluation": "认真负责",
                },
                "evaluation": {
                    "completeness_score": 75,
                    "strengths": ["信息完整"],
                    "weaknesses": ["经历不足"],
                    "missing_sections": ["教育经历", "工作经历"],
                },
            }

        final_state = run_resume_analysis_agent(state, llm=fake_llm)

        assert final_state["current_stage"] == "done"
        assert final_state["structured_resume"]["basic_info"]["name"] == "赵六"
        assert (
            final_state["structured_resume"]["basic_info"]["email"]
            == "zhao@example.com"
        )
        assert final_state["resume_evaluation"]["completeness_score"] == 75

    def test_agent_class_directly(self, tmp_path):
        """直接使用 ResumeAnalysisAgent 类。"""
        docx = tmp_path / "resume.docx"
        make_docx(docx, ["姓名：钱七", "电话：13600136000", "技能：Java、Spring"])
        state = initial_workflow_state(task_id="task-1", file_path=str(docx))

        agent = ResumeAnalysisAgent(use_configured_llm=False)
        final_state = agent.run(state)

        assert final_state["current_stage"] == "done"
        assert final_state["structured_resume"]["basic_info"]["name"] == "钱七"


# ========================================================================
# WorkflowState 初始化测试
# ========================================================================


class TestWorkflowState:
    """WorkflowState 初始化和字段测试。"""

    def test_initial_state_fields(self):
        """初始状态包含所有必需字段。"""
        state = initial_workflow_state(task_id="test", file_path="/fake/path.docx")
        assert state["task_id"] == "test"
        assert state["current_stage"] == "upload"
        assert state["error"] is None
        assert state["file_path"] == "/fake/path.docx"
        assert state["file_type"] == "unknown"
        assert state["converted_file_path"] is None
        assert state["raw_text"] == ""
        assert state["structured_resume"] == {}
        assert state["resume_evaluation"] == {}

    def test_initial_state_agent2_fields(self):
        """初始状态包含 Agent 2 预留字段。"""
        state = initial_workflow_state(task_id="test", file_path="/fake/path.docx")
        assert "jd_source_type" in state
        assert "jd_raw_text" in state
        assert "job_requirements" in state
        assert "match_result" in state
        assert "gap_report" in state

    def test_initial_state_agent3_fields(self):
        """初始状态包含 Agent 3 预留字段。"""
        state = initial_workflow_state(task_id="test", file_path="/fake/path.docx")
        assert "optimized_resume" in state
        assert "diff_report" in state
        assert "output_file_path" in state
        assert "optimization_summary" in state


if __name__ == "__main__":
    import pytest

    raise SystemExit(pytest.main([__file__, "-v"]))
