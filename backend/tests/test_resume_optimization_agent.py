"""Agent 3 工具与编排测试。"""

from __future__ import annotations

import sys
import threading
import time
import hashlib
from pathlib import Path

import fitz
from docx import Document


BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.agent.resume_optimization_agent import (  # noqa: E402
    run_resume_optimization_agent,
)
from app.schema.workflow_state import initial_workflow_state  # noqa: E402
from app.tools.diff_generator import diff_generator  # noqa: E402
from app.tools.doc_generator import generate_resume_document  # noqa: E402
from app.tools.section_rewriter import (  # noqa: E402
    build_section_rewrite_prompt,
    section_rewriter,
)
from app.schemas.optimization import SectionRewriteRequest  # noqa: E402


def _resume() -> dict:
    return {
        "basic_info": {
            "name": "张三",
            "phone": "13800138000",
            "email": "zhang@example.com",
            "location": "上海",
        },
        "education": [
            {
                "school": "示例大学",
                "major": "计算机科学",
                "degree": "本科",
                "period": "2020-2024",
            }
        ],
        "work_experience": [
            {
                "company": "示例科技",
                "position": "后端工程师",
                "period": "2024-至今",
                "description": "负责接口开发",
            },
            {
                "company": "示例数据",
                "position": "实习生",
                "period": "2023-2024",
                "description": "参与数据处理",
            },
        ],
        "project_experience": [
            {
                "name": "招聘助手",
                "role": "开发者",
                "description": "使用 Python 开发简历分析功能",
            }
        ],
        "skills": ["Python", "FastAPI"],
        "self_evaluation": "学习能力强",
    }


def _state() -> dict:
    state = initial_workflow_state("optimization-test", "resume.docx")
    state["structured_resume"] = _resume()
    state["job_requirements"] = {
        "title": "AI 应用开发工程师",
        "skills": ["Python", "FastAPI", "LangChain"],
        "responsibilities": ["开发 Agent 应用"],
        "qualifications": ["熟悉 Python"],
    }
    state["match_result"] = {"missing_skills": ["LangChain"]}
    state["gap_report"] = {
        "sections": [
            {"section_type": "work_experience", "status": "weak"},
            {"section_type": "project_experience", "status": "weak"},
            {"section_type": "self_evaluation", "status": "covered"},
        ],
        "critical_gaps": ["工作和项目成果表达不够具体"],
    }
    return state


def test_section_rewriter_returns_requested_schema():
    request = SectionRewriteRequest(
        section_type="work_experience",
        original_content="负责接口开发",
        gap_report={"weak": ["work_experience"]},
        job_requirements={"skills": ["Python"]},
        job_keywords=["Python"],
    )

    def fake_llm(_prompt: str):
        return {
            "section_type": "ignored",
            "original_content": "ignored",
            "rewritten_content": "使用 Python 负责核心接口开发",
            "change_reason": "突出技术栈和职责",
            "changes": [{"type": "modified", "description": "明确 Python 技术栈"}],
        }

    result = section_rewriter(request, llm=fake_llm)

    assert result.section_type == "work_experience"
    assert result.original_content == "负责接口开发"
    assert result.rewritten_content == "使用 Python 负责核心接口开发"
    assert result.changes[0].type == "modified"


def test_section_rewriter_uses_section_specific_guidance():
    prompts = {
        section_type: build_section_rewrite_prompt(
            SectionRewriteRequest(
                section_type=section_type,
                original_content="原始内容",
            )
        )
        for section_type in (
            "work_experience",
            "project_experience",
            "self_evaluation",
            "skills",
        )
    }

    assert "STAR" in prompts["work_experience"]
    assert "技术决策" in prompts["project_experience"]
    assert "空泛" in prompts["self_evaluation"]
    assert "去重" in prompts["skills"]
    assert len(set(prompts.values())) == 4


def test_diff_generator_produces_highlight_spans():
    original = _resume()
    optimized = _resume()
    optimized["self_evaluation"] = "学习能力强，具备团队协作意识"

    report = diff_generator(original, optimized)

    section = next(
        item for item in report.sections if item.section_type == "self_evaluation"
    )
    assert section.changed is True
    assert any(span.type == "added" for span in section.spans)


def test_doc_generator_edits_source_resume_without_clearing_template(tmp_path):
    template_path = tmp_path / "template.docx"
    template = Document()
    template.sections[0].header.paragraphs[0].text = "自定义简历模板"
    template.add_paragraph("项目经历")
    template.add_paragraph("负责接口开发")
    template.save(template_path)

    output_path = tmp_path / "optimized.docx"
    result = generate_resume_document(
        _resume(),
        str(output_path),
        source_document_path=str(template_path),
        text_edits=[
            {
                "paragraph_index": 1,
                "original_text": "负责接口开发",
                "rewritten_text": "负责核心接口开发",
            }
        ],
    )

    generated = Document(result)
    body_text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert generated.sections[0].header.paragraphs[0].text == "自定义简历模板"
    assert "项目经历" in body_text
    assert "负责核心接口开发" in body_text
    assert "张三" not in body_text


def test_agent_rewrites_selected_sections_in_parallel(tmp_path):
    state = _state()
    active = 0
    maximum_active = 0
    lock = threading.Lock()

    def fake_llm(prompt: str):
        nonlocal active, maximum_active
        with lock:
            active += 1
            maximum_active = max(maximum_active, active)
        time.sleep(0.05)
        with lock:
            active -= 1

        if '"original_content": "Python、FastAPI"' in prompt:
            rewritten = "Python、FastAPI、LangChain"
        elif "使用 Python 开发简历分析功能" in prompt:
            rewritten = "使用 Python 开发简历分析功能，完善 Agent 工作流"
        else:
            rewritten = "使用 Python 完成接口开发并提升交付效率"
        return {
            "rewritten_content": rewritten,
            "change_reason": "针对目标岗位强化表达",
            "changes": [{"type": "modified", "description": "强化岗位相关性"}],
        }

    final_state = run_resume_optimization_agent(
        state,
        llm=fake_llm,
        max_workers=4,
        output_dir=str(tmp_path),
    )

    assert final_state["current_stage"] == "done"
    assert maximum_active >= 2
    assert (
        state["structured_resume"]["work_experience"][0]["description"]
        == "负责接口开发"
    )
    assert (
        "提升交付效率"
        in final_state["optimized_resume"]["work_experience"][0]["description"]
    )
    assert "LangChain" in final_state["optimized_resume"]["skills"]
    assert final_state["optimized_resume"]["self_evaluation"] == "学习能力强"
    assert Path(final_state["output_file_path"]).is_file()
    assert final_state["diff_report"]["sections"]
    assert final_state["optimization_summary"]["modified_count"] >= 1


def test_agent_edits_uploaded_docx_in_place_instead_of_rebuilding_it(tmp_path):
    source_path = tmp_path / "uploaded-resume.docx"
    source = Document()
    source.sections[0].header.paragraphs[0].text = "用户原始页眉"
    source.add_paragraph("用户原始姓名")
    heading = source.add_paragraph("项目经历")
    heading.runs[0].bold = True
    project_title = source.add_paragraph("招聘助手")
    project_title.runs[0].bold = True
    source.add_paragraph("使用 Python 开发简历分析功能。")
    source.save(source_path)

    state = _state()
    state["file_path"] = str(source_path)
    state["structured_resume"]["project_experience"][0][
        "description"
    ] = "使用 Python 开发简历分析功能。"
    state["match_result"] = {}
    state["gap_report"] = {
        "sections": [{"section_type": "project_experience", "status": "weak"}]
    }

    def fake_llm(_prompt: str):
        return {
            "rewritten_content": "使用 Python 完成简历分析模块开发。",
            "change_reason": "明确项目动作",
            "changes": [{"type": "modified", "description": "强化动作表达"}],
        }

    final_state = run_resume_optimization_agent(
        state,
        llm=fake_llm,
        output_dir=str(tmp_path),
    )

    assert final_state["current_stage"] == "done"
    generated = Document(final_state["output_file_path"])
    assert [paragraph.text for paragraph in generated.paragraphs] == [
        "用户原始姓名",
        "项目经历",
        "招聘助手",
        "使用 Python 完成简历分析模块开发。",
    ]
    assert generated.sections[0].header.paragraphs[0].text == "用户原始页眉"
    assert (
        final_state["optimized_resume"]["project_experience"][0]["description"]
        == "使用 Python 完成简历分析模块开发。"
    )


def test_agent_edits_uploaded_pdf_and_generates_layout_preserving_word(tmp_path):
    source_path = tmp_path / "uploaded-resume.pdf"
    source = fitz.open()
    page = source.new_page(width=595, height=842)
    page.insert_text((40, 45), "ORIGINAL RESUME", fontsize=20, fontname="hebo")
    page.insert_text((40, 90), "SELF EVALUATION", fontsize=15, fontname="hebo")
    page.insert_text((40, 115), "Builds reliable Python services.", fontsize=10)
    page.insert_text((40, 175), "PROJECT EXPERIENCE", fontsize=15, fontname="hebo")
    page.insert_text((40, 205), "Resume Assistant", fontsize=12, fontname="hebo")
    page.insert_text(
        (40, 230),
        "Implemented a Python service with FastAPI.",
        fontsize=10,
    )
    page.insert_text((40, 290), "SKILLS", fontsize=15, fontname="hebo")
    page.insert_text((40, 315), "1. Python, FastAPI", fontsize=10)
    source.save(source_path)
    source.close()
    original_digest = hashlib.sha256(source_path.read_bytes()).hexdigest()

    state = _state()
    state["file_path"] = str(source_path)
    state["structured_resume"]["work_experience"] = []
    state["structured_resume"]["project_experience"][0][
        "description"
    ] = "Implemented a Python service with FastAPI."
    state["structured_resume"]["self_evaluation"] = "Builds reliable Python services."
    state["structured_resume"]["skills"] = ["Python", "FastAPI"]
    state["match_result"] = {}
    state["gap_report"] = {
        "sections": [
            {"section_type": "work_experience", "status": "missing"},
            {"section_type": "project_experience", "status": "weak"},
            {"section_type": "self_evaluation", "status": "weak"},
            {"section_type": "skills", "status": "weak"},
        ]
    }
    rewritten_sections: set[str] = set()

    def fake_llm(prompt: str):
        if '"section_type": "project_experience",\n  "original_content":' in prompt:
            rewritten_sections.add("project_experience")
            content = "Built a Python service with FastAPI."
        elif '"section_type": "self_evaluation",\n  "original_content":' in prompt:
            rewritten_sections.add("self_evaluation")
            content = "Builds reliable Python services for production."
        else:
            rewritten_sections.add("skills")
            content = "Python、FastAPI"
        return {
            "rewritten_content": content,
            "change_reason": "Strengthen job relevance",
            "changes": [{"type": "modified", "description": "Clarify evidence"}],
        }

    final_state = run_resume_optimization_agent(
        state,
        llm=fake_llm,
        output_dir=str(tmp_path),
    )

    assert final_state["current_stage"] == "done", final_state["error"]
    assert rewritten_sections == {"project_experience", "self_evaluation", "skills"}
    assert hashlib.sha256(source_path.read_bytes()).hexdigest() == original_digest
    generated = Document(final_state["output_file_path"])
    assert len(generated.inline_shapes) == 1
    assert generated.sections[0].page_width.pt == 595
    assert generated.sections[0].page_height.pt == 842
    assert final_state["diff_report"]["sections"]


def test_agent_keeps_gap_targets_and_builds_missing_skills_from_resume_evidence(
    tmp_path,
):
    """缺少技能时仍优化项目，并只从简历已有事实生成技能。"""
    state = _state()
    state["structured_resume"]["skills"] = []
    state["gap_report"] = {"critical_gaps": ["项目经历描述需要优化，技能关键词缺失"]}
    seen_sections: set[str] = set()

    def fake_llm(prompt: str):
        if '"section_type": "skills"' in prompt:
            seen_sections.add("skills")
            assert "使用 Python 开发简历分析功能" in prompt
            return {
                "rewritten_content": "Python、FastAPI",
                "change_reason": "从项目经历提取已有技能",
                "changes": [{"type": "added", "description": "补充已有技能"}],
            }

        seen_sections.add("project_experience")
        return {
            "rewritten_content": "使用 Python 开发简历分析功能，突出实现过程",
            "change_reason": "强化项目表达",
            "changes": [{"type": "modified", "description": "强化项目描述"}],
        }

    final_state = run_resume_optimization_agent(
        state,
        llm=fake_llm,
        output_dir=str(tmp_path),
    )

    assert final_state["current_stage"] == "done"
    assert seen_sections == {"project_experience", "skills"}
    assert final_state["optimized_resume"]["skills"] == ["Python", "FastAPI"]
    assert (
        "突出实现过程"
        in final_state["optimized_resume"]["project_experience"][0]["description"]
    )


def test_agent_reports_error_when_selected_sections_have_no_optimizable_content(
    tmp_path,
):
    """选中的段落没有任何内容时不能伪装成优化成功。"""
    state = _state()
    state["structured_resume"]["work_experience"] = []
    state["structured_resume"]["project_experience"] = []
    state["structured_resume"]["skills"] = []
    state["structured_resume"]["self_evaluation"] = ""
    state["match_result"] = {}
    state["gap_report"] = {
        "sections": [{"section_type": "work_experience", "status": "weak"}]
    }

    final_state = run_resume_optimization_agent(
        state,
        llm=lambda _prompt: {},
        output_dir=str(tmp_path),
    )

    assert final_state["current_stage"] == "error"
    assert "no optimizable content" in final_state["error"].casefold()
    assert final_state["output_file_path"] is None
