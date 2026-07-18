"""Layout-preserving DOCX editing tests for Agent 3."""

from __future__ import annotations

import hashlib
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.tools.docx_template_editor import (  # noqa: E402
    DocumentTextEdit,
    apply_resume_text_edits,
    extract_resume_text_slots,
)


def _source_resume(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.55)
    section.header.paragraphs[0].text = "保留的页眉"

    title = document.add_paragraph()
    title.style = "Body Text"
    title_run = title.add_run("张三的原始简历")
    title_run.bold = True
    title_run.font.size = Pt(20)

    heading = document.add_paragraph(style="Body Text")
    heading_run = heading.add_run("项目经历")
    heading_run.bold = True
    heading_run.font.size = Pt(15)

    project_title = document.add_paragraph(style="Body Text")
    project_title_run = project_title.add_run("招聘助手系统")
    project_title_run.bold = True
    project_title_run.font.size = Pt(12)

    document.add_paragraph("项目描述", style="Body Text")
    document.add_paragraph("使用 Python 开发简历分析功能。", style="Body Text")
    document.add_paragraph("技术亮点", style="Body Text")
    document.add_paragraph("使用 FastAPI 提供后端接口。", style="Body Text")
    document.save(path)


def _part_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            item.filename: hashlib.sha256(archive.read(item.filename)).hexdigest()
            for item in archive.infolist()
            if not item.is_dir()
        }


def test_extract_resume_text_slots_skips_template_and_subsection_labels(tmp_path):
    source = tmp_path / "source.docx"
    _source_resume(source)

    slots = extract_resume_text_slots(source, {"project_experience"})

    assert [slot.original_text for slot in slots] == [
        "使用 Python 开发简历分析功能。",
        "使用 FastAPI 提供后端接口。",
    ]
    assert all(slot.section_type == "project_experience" for slot in slots)


def test_apply_resume_text_edits_preserves_all_non_document_package_parts(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "optimized.docx"
    _source_resume(source)
    slot = extract_resume_text_slots(source, {"project_experience"})[0]

    apply_resume_text_edits(
        source,
        output,
        [
            DocumentTextEdit(
                paragraph_index=slot.paragraph_index,
                original_text=slot.original_text,
                rewritten_text="使用 Python 完成简历分析模块开发。",
            )
        ],
    )

    original = Document(source)
    optimized = Document(output)
    assert optimized.paragraphs[0].text == original.paragraphs[0].text
    assert optimized.paragraphs[1].text == "项目经历"
    assert optimized.paragraphs[4].text == "使用 Python 完成简历分析模块开发。"
    assert optimized.sections[0].left_margin == original.sections[0].left_margin
    assert optimized.sections[0].right_margin == original.sections[0].right_margin
    assert optimized.sections[0].header.paragraphs[0].text == "保留的页眉"

    before = _part_hashes(source)
    after = _part_hashes(output)
    assert before.keys() == after.keys()
    assert {
        name: digest for name, digest in before.items() if name != "word/document.xml"
    } == {name: digest for name, digest in after.items() if name != "word/document.xml"}


def test_missing_skills_uses_cloned_source_styles_instead_of_rebuilding_document(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "optimized.docx"
    _source_resume(source)

    apply_resume_text_edits(
        source,
        output,
        [],
        missing_skills_text="Python、FastAPI",
    )

    document = Document(output)
    assert document.paragraphs[-2].text == "专业技能"
    assert document.paragraphs[-2].runs[0].bold is True
    assert document.paragraphs[-1].text == "Python、FastAPI"
    before = _part_hashes(source)
    after = _part_hashes(output)
    assert {
        name: digest for name, digest in before.items() if name != "word/document.xml"
    } == {name: digest for name, digest in after.items() if name != "word/document.xml"}


def test_missing_skills_fills_existing_empty_skills_heading(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "optimized.docx"
    _source_resume(source)
    document = Document(source)
    heading = document.add_paragraph("技能")
    heading.runs[0].bold = True
    document.save(source)

    apply_resume_text_edits(
        source,
        output,
        [],
        missing_skills_text="Python、FastAPI",
    )

    generated = Document(output)
    texts = [paragraph.text for paragraph in generated.paragraphs]
    heading_index = texts.index("技能")
    assert texts[heading_index + 1] == "Python、FastAPI"
    assert texts.count("技能") == 1
