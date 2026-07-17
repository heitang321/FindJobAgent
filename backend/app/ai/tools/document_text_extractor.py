"""Tool 1.3: 文档文本提取器。

确定性操作，不调用 LLM。docx 使用 python-docx，pdf 使用 PyMuPDF（fitz）。
输出 {"raw_text": str, "char_count": int}。
"""
from __future__ import annotations

from pathlib import Path


def _extract_docx_text(path: Path) -> str:
    """使用 python-docx 提取 DOCX 文本。

    逐段读取段落文本，同时遍历表格单元格内容，用 ``|`` 分隔同一行单元格。
    """
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []

    # 提取段落文本
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # 提取表格文本
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> str:
    """使用 PyMuPDF（fitz）提取 PDF 文本。

    逐页读取文本内容，跳过空页面。
    """
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(str(path)) as document:
        for page in document:
            text = page.get_text().strip()
            if text:
                parts.append(text)
    return "\n".join(parts)


def document_text_extractor(file_path: str) -> dict[str, object]:
    """从简历文件中提取原始文本。

    根据文件扩展名选择提取方式：

    - ``.docx`` → python-docx
    - ``.pdf`` → PyMuPDF（fitz）

    Parameters
    ----------
    file_path:
        简历文件路径（通常是经过 pdf_to_word_converter 处理后的路径）。

    Returns
    -------
    dict[str, object]
        包含 ``raw_text``（提取的文本 str）和 ``char_count``（字符数 int）。
        对于不支持的文件类型，返回空文本。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        raw_text = _extract_docx_text(path)
    elif suffix == ".pdf":
        raw_text = _extract_pdf_text(path)
    else:
        raw_text = ""

    return {"raw_text": raw_text, "char_count": len(raw_text)}
"""Tool 1.3: extract text from PDF or Word resumes."""
from pathlib import Path


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> str:
    import fitz

    parts: list[str] = []
    with fitz.open(str(path)) as document:
        for page in document:
            text = page.get_text().strip()
            if text:
                parts.append(text)
    return "\n".join(parts)


def document_text_extractor(file_path: str) -> dict[str, object]:
    """Extract raw text and character count from a supported document."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        raw_text = _extract_docx_text(path)
    elif suffix == ".pdf":
        raw_text = _extract_pdf_text(path)
    else:
        raw_text = ""

    return {"raw_text": raw_text, "char_count": len(raw_text)}
