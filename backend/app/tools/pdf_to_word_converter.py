"""Tool 1.2：将 PDF 简历转换为保留版式的 Word 参考文件。

不会将 PDF 文本重建为可编辑的 Word 段落，因为这可能打乱内容顺序并破坏原模板。
每一页 PDF 都只渲染一次，并作为整页图片嵌入，从视觉上精确保留源文件。
Agent 分析阶段会单独读取原始 PDF。
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


_PDF_RENDER_SCALE = 2


def _build_layout_preserving_docx(pdf_path: Path, output_path: Path) -> None:
    """将渲染后的 PDF 页面嵌入 Word 文档，避免文本重排。"""
    import fitz
    from docx import Document
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()

    with fitz.open(str(pdf_path)) as pdf:
        if pdf.page_count == 0:
            raise ValueError("PDF contains no pages")

        first_page = pdf[0]
        section = document.sections[0]
        section.page_width = Pt(first_page.rect.width)
        section.page_height = Pt(first_page.rect.height)
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)

        matrix = fitz.Matrix(_PDF_RENDER_SCALE, _PDF_RENDER_SCALE)
        for index, page in enumerate(pdf):
            if index:
                run.add_break(WD_BREAK.PAGE)
            page_png = page.get_pixmap(matrix=matrix, alpha=False).tobytes("png")
            run.add_picture(
                BytesIO(page_png),
                width=Pt(page.rect.width),
                height=Pt(page.rect.height),
            )

    document.save(output_path)


def pdf_to_layout_preserving_docx(
    pdf_path: str | Path,
    output_path: str | Path,
) -> str:
    """将 PDF 页面嵌入到指定名称的 Word 输出文件。"""
    source = Path(pdf_path).resolve()
    destination = Path(output_path).resolve()
    if source.suffix.casefold() != ".pdf" or not source.is_file():
        raise ValueError("A valid PDF source is required.")
    if destination.suffix.casefold() != ".docx":
        raise ValueError("The layout-preserving output must be a DOCX file.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    _build_layout_preserving_docx(source, destination)
    return str(destination)


def pdf_to_word_converter(file_path: str) -> dict[str, object]:
    """将 PDF 转为视觉一致的 DOCX；DOCX 输入则保持不变。

    PDF 转换结果用于保留源模板的只读参考文件。它有意采用图片形式，
    不应作为 Agent 3 可编辑优化输出直接使用。
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return {"converted_path": str(path), "success": True}
    if suffix != ".pdf":
        return {"converted_path": str(path), "success": False}

    converted_path = path.with_suffix(".docx")
    try:
        _build_layout_preserving_docx(path, converted_path)
    except Exception:
        return {"converted_path": str(converted_path), "success": False}

    return {"converted_path": str(converted_path), "success": converted_path.exists()}
