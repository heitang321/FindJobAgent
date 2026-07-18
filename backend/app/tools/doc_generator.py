"""Tool 3.3: generate Agent 3's editable optimized-resume DOCX."""

from __future__ import annotations

from pathlib import Path
import tempfile
from typing import Any, Iterable

from app.tools.docx_template_editor import DocumentTextEdit, apply_resume_text_edits
from app.tools.pdf_template_editor import PdfTextEdit, apply_resume_pdf_edits
from app.tools.pdf_to_word_converter import pdf_to_layout_preserving_docx


_FONT_LATIN = "Calibri"
_FONT_CJK = "Microsoft YaHei"
_ACCENT = "2E74B5"


def _set_run_font(run, *, size=None, bold=None, color=None) -> None:
    from docx.oxml.ns import qn

    run.font.name = _FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _FONT_CJK)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _configure_document(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = _FONT_LATIN
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = _FONT_LATIN
    heading.font.size = Pt(13)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(_ACCENT)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(7)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_resume_header(document, basic_info: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    name = str(basic_info.get("name") or "优化简历")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(title.add_run(name), size=Pt(22), bold=True, color=RGBColor(0, 0, 0))

    contact_values = [
        str(basic_info.get(key) or "") for key in ("phone", "email", "location")
    ]
    contact_text = " | ".join(value for value in contact_values if value)
    if contact_text:
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(10)
        _set_run_font(
            contact.add_run(contact_text),
            size=Pt(10),
            color=RGBColor(90, 90, 90),
        )


def _add_heading(document, text: str) -> None:
    document.add_paragraph(text, style="Heading 1")


def _add_entry_title(document, values: list[str]) -> None:
    from docx.shared import Pt

    visible = [value for value in values if value]
    if not visible:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    _set_run_font(paragraph.add_run(" | ".join(visible)), bold=True)


def _description_lines(description: str) -> list[str]:
    lines = [
        line.strip().lstrip("-•· ")
        for line in description.splitlines()
        if line.strip().lstrip("-•· ")
    ]
    return lines or ([description.strip()] if description.strip() else [])


def _add_description(document, description: str) -> None:
    lines = _description_lines(description)
    for line in lines:
        paragraph = document.add_paragraph(
            style="List Bullet" if len(lines) > 1 else None
        )
        _set_run_font(paragraph.add_run(line))


def _add_education(document, education: list[dict[str, Any]]) -> None:
    if not education:
        return
    _add_heading(document, "教育经历")
    for item in education:
        _add_entry_title(
            document,
            [
                str(item.get("school") or ""),
                str(item.get("major") or ""),
                str(item.get("degree") or ""),
                str(item.get("period") or ""),
            ],
        )


def _add_experience_sections(document, resume: dict[str, Any]) -> None:
    work = resume.get("work_experience") or []
    if work:
        _add_heading(document, "工作经历")
        for item in work:
            _add_entry_title(
                document,
                [
                    str(item.get("company") or ""),
                    str(item.get("position") or ""),
                    str(item.get("period") or ""),
                ],
            )
            _add_description(document, str(item.get("description") or ""))

    projects = resume.get("project_experience") or []
    if projects:
        _add_heading(document, "项目经历")
        for item in projects:
            _add_entry_title(
                document,
                [str(item.get("name") or ""), str(item.get("role") or "")],
            )
            _add_description(document, str(item.get("description") or ""))


def generate_resume_document(
    optimized_resume: dict[str, Any],
    output_path: str,
    template_path: str | None = None,
    *,
    source_document_path: str | None = None,
    text_edits: Iterable[DocumentTextEdit | dict[str, Any]] = (),
    pdf_text_edits: Iterable[PdfTextEdit] = (),
    missing_skills_text: str = "",
) -> str:
    """Generate an editable Word resume and return its absolute path.

    When a source document is supplied, edit verified text slots in a copy of
    that DOCX. The source's layout and every package part except the body XML
    remain untouched. ``template_path`` is retained as a legacy alias for the
    source document path.
    """
    from docx import Document

    destination = Path(output_path).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)

    source = source_document_path or template_path
    if source:
        source_path = Path(source).resolve()
        if source_path.suffix.casefold() == ".pdf":
            with tempfile.NamedTemporaryFile(
                dir=destination.parent,
                prefix=f".{destination.stem}-",
                suffix=".pdf",
                delete=False,
            ) as temporary_file:
                temporary_pdf = Path(temporary_file.name)
            try:
                apply_resume_pdf_edits(source_path, temporary_pdf, pdf_text_edits)
                return pdf_to_layout_preserving_docx(temporary_pdf, destination)
            finally:
                if temporary_pdf.exists():
                    temporary_pdf.unlink()
        return apply_resume_text_edits(
            source_path,
            destination,
            text_edits,
            missing_skills_text=missing_skills_text,
        )

    document = Document()
    _configure_document(document)
    _add_resume_header(document, optimized_resume.get("basic_info") or {})
    _add_education(document, optimized_resume.get("education") or [])
    _add_experience_sections(document, optimized_resume)

    skills = optimized_resume.get("skills") or []
    if skills:
        _add_heading(document, "专业技能")
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run("、".join(str(skill) for skill in skills)))

    self_evaluation = str(optimized_resume.get("self_evaluation") or "")
    if self_evaluation:
        _add_heading(document, "自我评价")
        _add_description(document, self_evaluation)

    document.save(destination)
    return str(destination)
